# main.py — Part 1/3
# Full updated app with improved UX, hypothesis enrichment, tabbed PRD preview/edit,
# per-section regeneration, and a floating tips panel (manual refresh).
#
# NOTE: This file expects prompt_engine to expose:
#   - generate_hypotheses
#   - generate_hypothesis_details
#   - generate_experiment_plan (alias generate_prd may exist)
#   - validate_experiment_plan (optional)
#   - generate_dynamic_tips (optional; fallback provided)
#
# The UI only triggers LLM calls on explicit user actions (buttons) to avoid excess calls.

import json
import math
import random
import string
import textwrap
from io import BytesIO
from typing import Any, Dict, List, Optional

import streamlit as st

# Try optional libs for exports
PDF_AVAILABLE = False
DOCX_AVAILABLE = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Import prompt engine functions (best-effort)
PROMPT_ENGINE_AVAILABLE = False
_generate_hypotheses = None
_generate_hypothesis_details = None
_generate_experiment_plan = None
_validate_experiment_plan = None
_generate_dynamic_tips = None

try:
    from prompt_engine import (
        generate_hypotheses,
        generate_hypothesis_details,
        generate_experiment_plan,
        validate_experiment_plan,
    )

    # We'll map to local names and mark available
    _generate_hypotheses = generate_hypotheses
    _generate_hypothesis_details = generate_hypothesis_details
    _generate_experiment_plan = generate_experiment_plan
    _validate_experiment_plan = validate_experiment_plan if "validate_experiment_plan" in globals() else None
    # Try dynamic tips (optional)
    try:
        from prompt_engine import generate_dynamic_tips

        _generate_dynamic_tips = generate_dynamic_tips
    except Exception:
        _generate_dynamic_tips = None

    PROMPT_ENGINE_AVAILABLE = True
except Exception:
    PROMPT_ENGINE_AVAILABLE = False


# -------------------------
# Styling and UI helpers
# -------------------------
def inject_global_css():
    css = r"""
    <style>
    /* Page container */
    .block-container { max-width: 1200px; padding-top: 1rem; }

    /* Header gradient */
    header [role="banner"] { background: linear-gradient(90deg, #0ea5e9 0%, #6366f1 100%) !important; color: white !important; }

    /* Buttons and cards */
    .stButton > button { border-radius: 10px; padding: 0.5rem 0.75rem; font-weight: 600; }
    .card { border:1px solid #e6eef8; border-radius:12px; padding:1rem; background:linear-gradient(180deg,#ffffff,#fbfdff); box-shadow: 0 1px 3px rgba(15,23,42,0.04); }
    .muted { color: #6b7280; font-size:0.9rem; }

    /* Floating tips button */
    .floating-tips-button {
        position: fixed;
        right: 22px;
        bottom: 22px;
        z-index: 9999;
    }
    .floating-tips-panel {
        position: fixed;
        right: 22px;
        bottom: 80px;
        z-index: 9998;
        width: 360px;
        max-height: 70vh;
        overflow: auto;
        background: #ffffff;
        border-radius: 12px;
        box-shadow: 0 8px 30px rgba(2,6,23,0.2);
        border: 1px solid #e6eef8;
        padding: 16px;
    }
    .tip-card {
        background: #f8fafc;
        border-radius: 8px;
        padding: 10px;
        margin-bottom: 8px;
        font-size: 0.95rem;
    }

    /* Tab styling - make tabs look pill-like */
    .css-1avcm0n > div[role="tablist"] button {
        padding: 0.5rem 0.9rem;
        border-radius: 999px;
        margin-right: 6px;
    }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)


# -------------------------
# Utility functions
# -------------------------
def sanitize_text(x: Any) -> str:
    if x is None:
        return ""
    if isinstance(x, (dict, list)):
        try:
            return json.dumps(x, ensure_ascii=False)
        except Exception:
            return str(x)
    return str(x).strip()


def ensure_list(x: Any) -> List[Any]:
    if x is None:
        return []
    if isinstance(x, list):
        return x
    return [x]


def ensure_dict(x: Any) -> Dict[str, Any]:
    if isinstance(x, dict):
        return x
    return {}


def safe_int(x: Any, default: int = 0) -> int:
    try:
        if x is None:
            return default
        if isinstance(x, int):
            return x
        if isinstance(x, float):
            return int(x)
        s = str(x).strip()
        if s == "":
            return default
        return int(float(s))
    except Exception:
        return default


def safe_float(x: Any, default: float = 0.0) -> float:
    try:
        if x is None:
            return default
        if isinstance(x, (int, float)):
            return float(x)
        s = str(x).strip().replace("%", "")
        if s == "":
            return default
        return float(s)
    except Exception:
        return default


def extract_json_from_text(text: Optional[str]) -> Dict[str, Any]:
    """
    Extract JSON object from LLM messy output (best-effort). Returns dict or {}.
    """
    if not text:
        return {}
    t = text.strip()
    # direct parse
    try:
        return json.loads(t)
    except Exception:
        pass
    # find JSON block
    start = t.find("{")
    end = t.rfind("}")
    if start != -1 and end != -1 and end > start:
        candidate = t[start:end + 1]
        try:
            return json.loads(candidate)
        except Exception:
            pass
    # try list
    s = t.find("[")
    e = t.rfind("]")
    if s != -1 and e != -1 and e > s:
        candidate = t[s:e + 1]
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, list):
                return {"list": parsed}
        except Exception:
            pass
    return {}


def generate_experiment_id(prefix: str = "EXP") -> str:
    rand = "".join(random.choices(string.ascii_uppercase + string.digits, k=5))
    return f"{prefix}-{rand}"


# -------------------------
# Export helpers (PDF/DOCX/JSON)
# -------------------------
def pdf_safe(text: str) -> str:
    return text.replace("\t", "    ")


def generate_pdf_bytes_from_prd_dict(plan: Dict[str, Any]) -> Optional[bytes]:
    """
    Use ReportLab to generate a readable PDF. Returns bytes or None.
    """
    if not PDF_AVAILABLE:
        return None

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    def add_heading(txt, level=1):
        style = styles["Heading1"] if level == 1 else styles["Heading2"]
        story.append(Paragraph(pdf_safe(txt), style))
        story.append(Spacer(1, 8))

    def add_paragraph(txt):
        story.append(Paragraph(pdf_safe(txt), styles["BodyText"]))
        story.append(Spacer(1, 6))

    # Metadata
    meta = plan.get("metadata", {})
    title = meta.get("title", plan.get("metadata", {}).get("title", "Experiment PRD"))
    add_heading(title, level=1)
    add_paragraph(f"<b>Experiment ID:</b> {meta.get('experiment_id','')}")
    add_paragraph(f"<b>Team:</b> {meta.get('team','')}")
    add_paragraph(f"<b>Owner:</b> {meta.get('owner','')}")

    # Problem statement
    add_heading("Problem Statement", level=2)
    add_paragraph(plan.get("problem_statement", ""))

    # Hypotheses
    hyps = ensure_list(plan.get("hypotheses"))
    for i, h in enumerate(hyps, 1):
        add_heading(f"Hypothesis {i}", level=2)
        add_paragraph(f"<b>Hypothesis:</b> {h.get('hypothesis','')}")
        if h.get("rationale"):
            add_paragraph(f"<b>Rationale:</b> {h.get('rationale')}")
        if h.get("example_implementation"):
            add_paragraph(f"<b>Example Implementation:</b> {h.get('example_implementation')}")
        if h.get("behavioral_basis"):
            add_paragraph(f"<b>Behavioral Basis:</b> {h.get('behavioral_basis')}")

    # Proposed solution & variants
    add_heading("Proposed Solution & Variants", level=2)
    add_paragraph(plan.get("proposed_solution", ""))
    vars_table = [["Control", "Variation", "Notes"]]
    for v in ensure_list(plan.get("variants")):
        vars_table.append([v.get("control", ""), v.get("variation", ""), v.get("notes", "")])
    t = Table(vars_table, colWidths=[150, 150, 150])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    story.append(t)
    story.append(Spacer(1, 8))

    # Metrics
    metrics_table = [["Metric", "Formula", "Importance"]]
    for m in ensure_list(plan.get("metrics")):
        metrics_table.append([m.get("name", ""), m.get("formula", ""), m.get("importance", "")])
    t2 = Table(metrics_table, colWidths=[150, 220, 80])
    t2.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey)]))
    story.append(t2)
    story.append(Spacer(1, 8))

    # Risks & other sections
    add_heading("Risks & Mitigation", level=2)
    for r in ensure_list(plan.get("risks_and_assumptions")):
        add_paragraph(f"{r.get('risk','')} — {r.get('mitigation','')} (Severity: {r.get('severity','')})")

    add_heading("Experiment Design & Rollout Plan", level=2)
    ed = ensure_dict(plan.get("experiment_design", {}))
    add_paragraph(f"Traffic Allocation: {ed.get('traffic_allocation','')}")
    add_paragraph(f"Sample size / variant: {ed.get('sample_size_per_variant','')}")
    add_paragraph(f"Total sample size: {ed.get('total_sample_size','')}")
    add_paragraph(f"Estimated Duration (days): {ed.get('test_duration_days','')}")

    # Success & Learning
    add_heading("Success & Learning Criteria", level=2)
    sc = ensure_dict(plan.get("success_criteria", {}))
    add_paragraph(f"Confidence level: {sc.get('confidence_level','')}%")
    add_paragraph(f"Power: {sc.get('power','')}%")
    add_paragraph(f"MDE: {sc.get('MDE','')}%")

    # Statistical rationale
    if plan.get("statistical_rationale"):
        add_heading("Statistical Rationale", level=2)
        add_paragraph(plan.get("statistical_rationale"))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def generate_docx_bytes_from_plan(plan: Dict[str, Any]) -> Optional[bytes]:
    """
    Create a DOCX file bytes if python-docx is available.
    """
    if not DOCX_AVAILABLE:
        return None

    doc = Document()
    meta = plan.get("metadata", {})
    doc.add_heading(meta.get("title", "Experiment PRD"), level=0)
    doc.add_paragraph(f"Experiment ID: {meta.get('experiment_id','')}")
    doc.add_paragraph(f"Team: {meta.get('team','')}")
    doc.add_paragraph(f"Owner: {meta.get('owner','')}")

    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph(plan.get("problem_statement", ""))

    hyps = ensure_list(plan.get("hypotheses"))
    for i, h in enumerate(hyps, 1):
        doc.add_heading(f"Hypothesis {i}", level=1)
        doc.add_paragraph(h.get("hypothesis", ""))
        if h.get("rationale"):
            doc.add_paragraph("Rationale: " + h.get("rationale", ""))
        if h.get("example_implementation"):
            doc.add_paragraph("Example Implementation: " + h.get("example_implementation", ""))

    doc.add_heading("Proposed Solution & Variants", level=1)
    doc.add_paragraph(plan.get("proposed_solution", ""))

    # Finish file
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()


# -------------------------
# Plan normalization / defaults
# -------------------------
DEFAULT_PLAN: Dict[str, Any] = {
    "metadata": {"title": "Untitled Experiment", "team": "", "owner": "", "experiment_id": ""},
    "problem_statement": "",
    "hypotheses": [],
    "proposed_solution": "",
    "variants": [{"control": "", "variation": "", "notes": ""}],
    "metrics": [{"name": "", "formula": "", "importance": "Primary"}],
    "guardrail_metrics": [{"name": "", "direction": "Decrease", "threshold": ""}],
    "experiment_design": {
        "traffic_allocation": "50/50",
        "sample_size_per_variant": 0,
        "total_sample_size": 0,
        "test_duration_days": 0,
        "dau_coverage_percent": 0.0,
        "power": 80.0,
    },
    "success_criteria": {"confidence_level": 95.0, "power": 80.0, "MDE": 1.0, "benchmark": "", "monitoring": ""},
    "success_learning_criteria": {"definition_of_success": "", "stopping_rules": "", "rollback_criteria": ""},
    "risks_and_assumptions": [{"risk": "", "severity": "Medium", "mitigation": ""}],
    "statistical_rationale": "",
}


def sanitize_plan(plan: Dict[str, Any]) -> Dict[str, Any]:
    """
    Normalize shapes, fill defaults, coerce types.
    """
    if not isinstance(plan, dict):
        plan = {}
    merged = json.loads(json.dumps(DEFAULT_PLAN))  # deep copy

    # metadata
    merged["metadata"].update(ensure_dict(plan.get("metadata", {})))
    if not merged["metadata"].get("experiment_id"):
        merged["metadata"]["experiment_id"] = generate_experiment_id()

    # scalars
    for k in ["problem_statement", "proposed_solution", "statistical_rationale"]:
        if k in plan:
            merged[k] = sanitize_text(plan.get(k))

    # hypotheses
    merged["hypotheses"] = []
    for h in ensure_list(plan.get("hypotheses")):
        if isinstance(h, dict):
            merged["hypotheses"].append(
                {
                    "hypothesis": sanitize_text(h.get("hypothesis", "")),
                    "rationale": sanitize_text(h.get("rationale", "")),
                    "example_implementation": sanitize_text(h.get("example_implementation", "")),
                    "behavioral_basis": sanitize_text(h.get("behavioral_basis", "")),
                }
            )

    # variants
    merged["variants"] = []
    for v in ensure_list(plan.get("variants")):
        if isinstance(v, dict):
            merged["variants"].append(
                {"control": sanitize_text(v.get("control", "")), "variation": sanitize_text(v.get("variation", "")), "notes": sanitize_text(v.get("notes", ""))}
            )

    # metrics
    merged["metrics"] = []
    for m in ensure_list(plan.get("metrics")):
        if isinstance(m, dict):
            importance = m.get("importance", "Primary")
            if importance not in ("Primary", "Secondary"):
                importance = "Primary"
            merged["metrics"].append({"name": sanitize_text(m.get("name", "")), "formula": sanitize_text(m.get("formula", "")), "importance": importance})

    # guardrails
    merged["guardrail_metrics"] = []
    for g in ensure_list(plan.get("guardrail_metrics")):
        if isinstance(g, dict):
            direction = g.get("direction", "Decrease")
            if direction not in ("Increase", "Decrease", "No Change"):
                direction = "Decrease"
            merged["guardrail_metrics"].append({"name": sanitize_text(g.get("name", "")), "direction": direction, "threshold": sanitize_text(g.get("threshold", ""))})

    # experiment design
    ed = ensure_dict(plan.get("experiment_design", {}))
    merged["experiment_design"] = {
        "traffic_allocation": sanitize_text(ed.get("traffic_allocation", "50/50")),
        "sample_size_per_variant": safe_int(ed.get("sample_size_per_variant", 0)),
        "total_sample_size": safe_int(ed.get("total_sample_size", 0)),
        "test_duration_days": safe_int(ed.get("test_duration_days", 0)),
        "dau_coverage_percent": safe_float(ed.get("dau_coverage_percent", 0.0)),
        "power": safe_float(ed.get("power", 80.0)),
    }

    # success criteria
    sc = ensure_dict(plan.get("success_criteria", {}))
    merged["success_criteria"] = {
        "confidence_level": safe_float(sc.get("confidence_level", 95.0)),
        "power": safe_float(sc.get("power", merged["experiment_design"]["power"])),
        "MDE": safe_float(sc.get("MDE", 1.0)),
        "benchmark": sanitize_text(sc.get("benchmark", "")),
        "monitoring": sanitize_text(sc.get("monitoring", "")),
    }

    # success learning criteria
    sl = ensure_dict(plan.get("success_learning_criteria", {}))
    merged["success_learning_criteria"] = {
        "definition_of_success": sanitize_text(sl.get("definition_of_success", "")),
        "stopping_rules": sanitize_text(sl.get("stopping_rules", "")),
        "rollback_criteria": sanitize_text(sl.get("rollback_criteria", "")),
    }

    # risks
    merged["risks_and_assumptions"] = []
    for r in ensure_list(plan.get("risks_and_assumptions")):
        if isinstance(r, dict):
            sev = r.get("severity", "Medium")
            if sev not in ("High", "Medium", "Low"):
                sev = "Medium"
            merged["risks_and_assumptions"].append({"risk": sanitize_text(r.get("risk", "")), "severity": sev, "mitigation": sanitize_text(r.get("mitigation", ""))})

    merged["statistical_rationale"] = sanitize_text(plan.get("statistical_rationale", ""))

    return merged
# -------------------------
# Small helper: Render PRD as Markdown (for Preview)
# -------------------------
def prd_to_markdown(plan: Dict[str, Any]) -> str:
    plan = sanitize_plan(plan)
    md = []
    meta = plan.get("metadata", {})
    md.append(f"# {meta.get('title', 'Experiment PRD')}")
    md.append("")
    md.append(f"**Owner:** {meta.get('owner','-')}  |  **Team:** {meta.get('team','-')}  |  **ID:** {meta.get('experiment_id','-')}")
    md.append("")
    md.append("## Problem Statement")
    md.append(plan.get("problem_statement", "") or "-")
    md.append("")
    md.append("## Hypothesis")
    hyps = ensure_list(plan.get("hypotheses"))
    if hyps:
        for i, h in enumerate(hyps, 1):
            md.append(f"### Hypothesis {i}")
            md.append(h.get("hypothesis", ""))
            if h.get("rationale"):
                md.append(f"**Rationale:** {h.get('rationale')}")
            if h.get("example_implementation"):
                md.append(f"**Example Implementation:** {h.get('example_implementation')}")
            if h.get("behavioral_basis"):
                md.append(f"**Behavioral Basis:** {h.get('behavioral_basis')}")
            md.append("")
    else:
        md.append("- Not provided")
        md.append("")

    md.append("## Proposed Solution & Variants")
    md.append(plan.get("proposed_solution", "") or "-")
    md.append("")
    for v in ensure_list(plan.get("variants")):
        md.append(f"- **Control:** {v.get('control','')}")
        md.append(f"- **Variation:** {v.get('variation','')}")
        if v.get("notes"):
            md.append(f"  - _Notes:_ {v.get('notes')}")
    md.append("")

    if plan.get("metrics"):
        md.append("## Success Metrics")
        for m in plan.get("metrics", []):
            md.append(f"- **{m.get('name','')}** — {m.get('formula','')} (_{m.get('importance','')}_)")
        md.append("")

    if plan.get("guardrail_metrics"):
        md.append("## Guardrails")
        for g in plan.get("guardrail_metrics", []):
            md.append(f"- **{g.get('name','')}** — {g.get('direction','')} {g.get('threshold','')}")
        md.append("")

    md.append("## Experiment Design & Rollout")
    ed = ensure_dict(plan.get("experiment_design", {}))
    md.append(f"- **Traffic Allocation:** {ed.get('traffic_allocation','')}")
    md.append(f"- **Sample Size / Variant:** {ed.get('sample_size_per_variant','')}")
    md.append(f"- **Total Sample Size:** {ed.get('total_sample_size','')}")
    md.append(f"- **Estimated Duration (days):** {ed.get('test_duration_days','')}")
    if ed.get("dau_coverage_percent") is not None:
        md.append(f"- **DAU Coverage:** {ed.get('dau_coverage_percent')}%")
    md.append("")

    if plan.get("risks_and_assumptions"):
        md.append("## Risks & Mitigation")
        for r in plan.get("risks_and_assumptions", []):
            md.append(f"- **{r.get('risk','')}** (_{r.get('severity','')}_): {r.get('mitigation','')}")
        md.append("")

    sl = ensure_dict(plan.get("success_learning_criteria", {}))
    if any(sl.values()):
        md.append("## Success & Learning Criteria")
        md.append(f"- **Definition of Success:** {sl.get('definition_of_success','')}")
        md.append(f"- **Stopping Rules:** {sl.get('stopping_rules','')}")
        md.append(f"- **Rollback Criteria:** {sl.get('rollback_criteria','')}")
        md.append("")

    sc = ensure_dict(plan.get("success_criteria", {}))
    md.append("## Statistical Rationale & Success Criteria")
    md.append(f"- **Confidence Level:** {sc.get('confidence_level','')}%")
    md.append(f"- **Power:** {sc.get('power','')}%")
    md.append(f"- **MDE:** {sc.get('MDE','')}%")
    if plan.get("statistical_rationale"):
        md.append("")
        md.append("### Statistical Rationale")
        md.append(plan.get("statistical_rationale", ""))

    return "\n\n".join(md)


# -------------------------
# Tips helper (uses prompt engine if available)
# -------------------------
def generate_tips(context: Dict[str, Any], current_step: str) -> List[str]:
    """
    Use prompt_engine.generate_dynamic_tips if available, otherwise return static fallback tips.
    """
    # Try LLM
    try:
        if PROMPT_ENGINE_AVAILABLE and _generate_dynamic_tips:
            out = _generate_dynamic_tips(context, current_step)
            if isinstance(out, list) and out:
                return out
            # if string, try splitting by newlines
            if isinstance(out, str):
                lines = [l.strip() for l in out.splitlines() if l.strip()]
                return lines[:5]
    except Exception:
        pass

    # Static fallback tips (short & actionable)
    base_tips = {
        "inputs": [
            "🎯 Keep your business goal crisp and measurable (e.g., 'Increase checkout conversion by 3%').",
            "📊 Use a single, well-defined primary metric — avoid compound metrics.",
            "🔢 Provide baseline numbers (e.g., '25% CTR') so the tool can size experiments."
        ],
        "hypothesis": [
            "✍️ Make your hypothesis falsifiable: 'If X, then Y for Z users.'",
            "📌 Tie the hypothesis to the metric and the user persona explicitly.",
            "⚡ Start small — propose lightweight UI or copy changes first."
        ],
        "prd": [
            "🛡️ Add at least one guardrail metric to prevent regressions (e.g., retention, error rate).",
            "📐 Ensure sample size estimates are realistic given DAU coverage.",
            "🔁 Define stopping and rollback rules clearly to reduce risk during rollout."
        ]
    }
    return base_tips.get(current_step, base_tips["prd"])


# -------------------------
# Main app UI
# -------------------------
def main():
    st.set_page_config(page_title="PM Experiment Architect", page_icon="🧪", layout="wide")
    inject_global_css()

    # App title + header row
    header_col_left, header_col_right = st.columns([8, 2])
    with header_col_left:
        st.title("🧪 PM Experiment Architect")
        st.caption("AI co-pilot to generate high-quality, shareable A/B test PRDs. Work step-by-step.")
    with header_col_right:
        # tips toggle (top-right)
        if "show_tips_panel" not in st.session_state:
            st.session_state["show_tips_panel"] = False
        if st.session_state["show_tips_panel"]:
            if st.button("Close Tips ❌"):
                st.session_state["show_tips_panel"] = False
        else:
            if st.button("💡 Tips"):
                st.session_state["show_tips_panel"] = True

    # -------------------------
    # Sidebar — Step 1: Inputs
    # -------------------------
    with st.sidebar:
        st.header("Step 1 — Inputs")
        st.markdown("Enter the minimal context needed for a focused experiment PRD.")
        st.session_state["high_level_goal"] = st.text_input(
            "High-Level Business Goal",
            value=st.session_state.get("high_level_goal", ""),
            placeholder="e.g., Increase checkout conversion rate"
        )
        st.session_state["product_type"] = st.selectbox(
            "Product Type",
            ["Mobile App", "E-commerce Website", "SaaS Dashboard", "Marketplace", "B2B Admin Console"],
            index=0 if "product_type" not in st.session_state else ["Mobile App", "E-commerce Website", "SaaS Dashboard", "Marketplace", "B2B Admin Console"].index(st.session_state.get("product_type"))
        )
        st.session_state["target_user"] = st.text_input(
            "Target User Persona",
            value=st.session_state.get("target_user", ""),
            placeholder="e.g., First-time visitor"
        )
        st.session_state["key_metric"] = st.text_input(
            "Primary Metric",
            value=st.session_state.get("key_metric", ""),
            placeholder="e.g., CTR, Purchase Conversion Rate"
        )

        col1, col2 = st.columns(2)
        with col1:
            st.session_state["current_value"] = st.text_input("Current Value (baseline)", value=st.session_state.get("current_value", ""), placeholder="e.g., 25% or 0.25")
        with col2:
            st.session_state["target_value"] = st.text_input("Target Value", value=st.session_state.get("target_value", ""), placeholder="e.g., 28% or 0.28")

        st.markdown("---")
        st.subheader("Optional: Statistical Defaults")
        st.session_state["confidence_level"] = st.slider("Confidence Level (%)", min_value=80, max_value=99, value=int(st.session_state.get("confidence_level", 95)), step=1)
        st.session_state["power"] = st.slider("Statistical Power (%)", min_value=70, max_value=99, value=int(st.session_state.get("power", 80)), step=1)
        st.session_state["MDE"] = st.slider("Minimum Detectable Effect (%)", min_value=0, max_value=50, value=int(st.session_state.get("MDE", 1)), step=1)

        st.markdown("---")
        st.subheader("Optional: Experiment Design")
        st.session_state["dau_coverage_percent"] = st.slider("DAU Coverage (%)", min_value=0, max_value=100, value=int(st.session_state.get("dau_coverage_percent", 0)), step=1)
        st.session_state["traffic_allocation"] = st.selectbox("Traffic Allocation", ["50/50", "90/10", "80/20", "70/30", "60/40"], index=0 if "traffic_allocation" not in st.session_state else ["50/50", "90/10", "80/20", "70/30", "60/40"].index(st.session_state.get("traffic_allocation", "50/50")))

        # Save compact context object for easy passing
        st.session_state["sidebar_context"] = {
            "high_level_goal": st.session_state.get("high_level_goal", ""),
            "product_type": st.session_state.get("product_type", ""),
            "target_user": st.session_state.get("target_user", ""),
            "key_metric": st.session_state.get("key_metric", ""),
            "current_value": st.session_state.get("current_value", ""),
            "target_value": st.session_state.get("target_value", ""),
            "success_criteria": {"confidence_level": st.session_state.get("confidence_level", 95), "power": st.session_state.get("power", 80), "MDE": st.session_state.get("MDE", 1)},
            "experiment_design": {"dau_coverage_percent": st.session_state.get("dau_coverage_percent", 0), "traffic_allocation": st.session_state.get("traffic_allocation", "50/50")},
            "metadata": {"owner": st.session_state.get("owner", ""), "team": st.session_state.get("team", ""), "experiment_id": st.session_state.get("experiment_id", generate_experiment_id())}
        }

    # -------------------------
    # Main: Step 2 — Hypothesis Entry & Enrichment
    # -------------------------
    st.header("Step 2 — Hypothesis")
    st.markdown("Write a single hypothesis and click **Enrich Hypothesis**. We'll auto-generate rationale, example implementation, and behavioral basis — you can edit them after.")

    user_hyp = st.text_area("Your Hypothesis (one or two lines)", value=st.session_state.get("user_hypothesis_text", ""), height=120, key="user_hypothesis_text")
    col_a, col_b = st.columns([1, 1])
    with col_a:
        enrich_clicked = st.button("Enrich Hypothesis with AI", use_container_width=True)
    with col_b:
        clear_hyp = st.button("Clear Hypothesis", use_container_width=True)

    if clear_hyp:
        for k in ["chosen_hypothesis", "edited_hypothesis", "edited_rationale", "edited_example", "edited_behavioral"]:
            if k in st.session_state:
                del st.session_state[k]

    if enrich_clicked:
        if not user_hyp or not user_hyp.strip():
            st.warning("Please enter a hypothesis first.")
        else:
            with st.spinner("Expanding hypothesis..."):
                try:
                    ctx = st.session_state.get("sidebar_context", {})
                    if PROMPT_ENGINE_AVAILABLE and _generate_hypothesis_details:
                        raw = _generate_hypothesis_details(user_hyp.strip(), ctx)
                        if isinstance(raw, str):
                            parsed = extract_json_from_text(raw)
                        elif isinstance(raw, dict):
                            parsed = raw
                        else:
                            parsed = {}
                    else:
                        # Local fallback: minimal expansion
                        parsed = {
                            "hypothesis": user_hyp.strip(),
                            "rationale": "",
                            "example_implementation": "",
                            "behavioral_basis": ""
                        }

                    chosen = {
                        "hypothesis": sanitize_text(parsed.get("hypothesis", user_hyp.strip())),
                        "rationale": sanitize_text(parsed.get("rationale", "")),
                        "example_implementation": sanitize_text(parsed.get("example_implementation", "")),
                        "behavioral_basis": sanitize_text(parsed.get("behavioral_basis", ""))
                    }
                    st.session_state["chosen_hypothesis"] = chosen

                    # Pre-fill editable fields
                    st.session_state["edited_hypothesis"] = chosen["hypothesis"]
                    st.session_state["edited_rationale"] = chosen["rationale"]
                    st.session_state["edited_example"] = chosen["example_implementation"]
                    st.session_state["edited_behavioral"] = chosen["behavioral_basis"]

                    st.success("Hypothesis expanded. Review & edit the generated details below.")
                except Exception as e:
                    st.error(f"Failed to expand hypothesis: {e}")

    # Show editable generated hypothesis
    if "chosen_hypothesis" in st.session_state:
        st.subheader("Review & Edit Hypothesis")
        ch = st.session_state["chosen_hypothesis"]
        col1, col2 = st.columns(2)
        with col1:
            st.session_state["edited_hypothesis"] = st.text_area("Hypothesis", value=st.session_state.get("edited_hypothesis", ch.get("hypothesis", "")), height=100, key="edited_hypothesis")
            st.session_state["edited_behavioral"] = st.text_input("Behavioral Basis", value=st.session_state.get("edited_behavioral", ch.get("behavioral_basis", "")), key="edited_behavioral")
        with col2:
            st.session_state["edited_rationale"] = st.text_area("Rationale", value=st.session_state.get("edited_rationale", ch.get("rationale", "")), height=120, key="edited_rationale")
            st.session_state["edited_example"] = st.text_area("Example Implementation", value=st.session_state.get("edited_example", ch.get("example_implementation", "")), height=120, key="edited_example")

        # Keep chosen_hypothesis synced with edits
        st.session_state["chosen_hypothesis"] = {
            "hypothesis": sanitize_text(st.session_state.get("edited_hypothesis", "")),
            "rationale": sanitize_text(st.session_state.get("edited_rationale", "")),
            "example_implementation": sanitize_text(st.session_state.get("edited_example", "")),
            "behavioral_basis": sanitize_text(st.session_state.get("edited_behavioral", ""))
        }

    st.markdown("---")

    # -------------------------
    # Step 3 — Generate PRD
    # -------------------------
    st.header("Step 3 — Generate PRD")
    st.markdown("When ready, generate a fully personalized PRD. You can then edit sections individually or regenerate sections with AI.")

    col_left, col_right = st.columns([2, 1])
    with col_left:
        owner = st.text_input("Experiment Owner", value=st.session_state.get("owner", ""))
        team = st.text_input("Team", value=st.session_state.get("team", ""))
        title = st.text_input("Experiment Title", value=st.session_state.get("title", st.session_state.get("high_level_goal", "Untitled Experiment")))
    with col_right:
        st.markdown("**Preview** will update after generation. Use the tabs to switch between Edit and Preview.")

    generate_prd_clicked = st.button("Generate PRD", type="primary", use_container_width=True)
    if generate_prd_clicked:
        if "chosen_hypothesis" not in st.session_state:
            st.error("Please create and enrich a hypothesis first (Step 2).")
        else:
            with st.spinner("Generating PRD..."):
                try:
                    ctx = st.session_state.get("sidebar_context", {})
                    # Merge metadata
                    ctx["metadata"] = {"owner": owner, "team": team, "title": title, "experiment_id": st.session_state.get("experiment_id", generate_experiment_id())}
                    hyp = st.session_state["chosen_hypothesis"]

                    if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                        raw_plan = _generate_experiment_plan(ctx, hyp)
                        if isinstance(raw_plan, str):
                            parsed = extract_json_from_text(raw_plan)
                        elif isinstance(raw_plan, dict):
                            parsed = raw_plan
                        else:
                            parsed = {}
                    else:
                        # Fallback plan stub
                        parsed = {
                            "metadata": ctx.get("metadata", {}),
                            "problem_statement": f"Aim: {ctx.get('high_level_goal','')}. Move {ctx.get('key_metric','primary metric')} from {ctx.get('current_value','')} to {ctx.get('target_value','')}.",
                            "hypotheses": [hyp],
                            "proposed_solution": hyp.get("example_implementation", ""),
                            "variants": [{"control": "Current experience", "variation": hyp.get("example_implementation", ""), "notes": ""}],
                            "metrics": [{"name": ctx.get("key_metric", "Primary Metric"), "formula": "", "importance": "Primary"}],
                            "guardrail_metrics": [{"name": "Retention", "direction": "Decrease", "threshold": ""}],
                            "experiment_design": {"traffic_allocation": ctx.get("experiment_design", {}).get("traffic_allocation", "50/50"), "sample_size_per_variant": None, "total_sample_size": None, "test_duration_days": None, "dau_coverage_percent": ctx.get("experiment_design", {}).get("dau_coverage_percent", 0)},
                            "success_criteria": ctx.get("success_criteria", {}),
                            "risks_and_assumptions": [{"risk": "Novelty effect", "severity": "Medium", "mitigation": "Monitor over time"}],
                            "success_learning_criteria": {"definition_of_success": "", "stopping_rules": "", "rollback_criteria": ""},
                            "statistical_rationale": ""
                        }

                    plan = sanitize_plan(parsed)
                    # ensure metadata fields
                    plan["metadata"].update({"owner": owner, "team": team, "title": title})
                    st.session_state["experiment_plan"] = plan
                    st.success("PRD generated. Open Edit tab to refine sections or Preview to view the final document.")
                except Exception as e:
                    st.error(f"Error generating PRD: {e}")

    # If there's a plan, show tabs for Edit / Preview
    plan = st.session_state.get("experiment_plan")
    if plan:
        # Keep a working copy that reflects user edits instantly
        if "final_prd" not in st.session_state:
            st.session_state["final_prd"] = plan

        tab1, tab2 = st.tabs(["✏️ Edit", "👀 Preview"])
        with tab1:
            st.markdown("### Edit PRD (per-section). Use the small **Regenerate** buttons to AI-refresh a single section.")
            # Metadata editor
            with st.expander("📇 Metadata", expanded=False):
                meta = ensure_dict(st.session_state["final_prd"].get("metadata", {}))
                meta["title"] = st.text_input("Title", value=meta.get("title", ""), key="md_title")
                meta["team"] = st.text_input("Team", value=meta.get("team", ""), key="md_team")
                meta["owner"] = st.text_input("Owner", value=meta.get("owner", ""), key="md_owner")
                meta["experiment_id"] = st.text_input("Experiment ID", value=meta.get("experiment_id", ""), key="md_eid")
                st.session_state["final_prd"]["metadata"] = meta

            # Problem Statement
            with st.expander("🧭 Problem Statement", expanded=False):
                st.session_state["final_prd"]["problem_statement"] = st.text_area("Problem Statement", value=st.session_state["final_prd"].get("problem_statement", ""), height=120, key="edit_problem")
                if st.button("♻️ Regenerate Problem Statement", key="regen_problem"):
                    # regenerate full plan and replace problem_statement
                    with st.spinner("Regenerating problem statement..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("chosen_hypothesis", {})
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["problem_statement"] = new_plan.get("problem_statement", st.session_state["final_prd"].get("problem_statement", ""))
                                st.success("Problem statement regenerated.")
                            else:
                                st.warning("LLM not available; cannot regenerate.")
                        except Exception as e:
                            st.error(f"Failed to regenerate: {e}")

            # Hypothesis editor
            with st.expander("💡 Hypothesis", expanded=False):
                hyps = ensure_list(st.session_state["final_prd"].get("hypotheses", []))
                if not hyps:
                    hyps = [st.session_state.get("chosen_hypothesis", {"hypothesis":"", "rationale":"", "example_implementation":"", "behavioral_basis":""})]
                # show first hypothesis only (single-hypothesis flow)
                h0 = hyps[0]
                h0["hypothesis"] = st.text_area("Hypothesis", value=h0.get("hypothesis", ""), height=100, key="edit_hyp")
                h0["rationale"] = st.text_area("Rationale", value=h0.get("rationale", ""), height=120, key="edit_rat")
                h0["example_implementation"] = st.text_area("Example Implementation", value=h0.get("example_implementation", ""), height=120, key="edit_ex")
                h0["behavioral_basis"] = st.text_input("Behavioral Basis", value=h0.get("behavioral_basis", ""), key="edit_beh")
                st.session_state["final_prd"]["hypotheses"] = [h0]

                if st.button("♻️ Regenerate Hypothesis Details", key="regen_hypothesis"):
                    with st.spinner("Regenerating hypothesis details..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            # use the current hypothesis text as seed
                            seed_hyp = h0.get("hypothesis","")
                            if PROMPT_ENGINE_AVAILABLE and _generate_hypothesis_details:
                                raw = _generate_hypothesis_details(seed_hyp, ctx)
                                parsed = raw if isinstance(raw, dict) else extract_json_from_text(raw)
                                if isinstance(parsed, dict):
                                    updated = {
                                        "hypothesis": parsed.get("hypothesis", seed_hyp),
                                        "rationale": parsed.get("rationale", ""),
                                        "example_implementation": parsed.get("example_implementation", ""),
                                        "behavioral_basis": parsed.get("behavioral_basis", "")
                                    }
                                    st.session_state["final_prd"]["hypotheses"] = [updated]
                                    st.success("Hypothesis details regenerated.")
                                else:
                                    st.warning("Could not parse regenerated hypothesis details.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate hypothesis: {e}")

            # Proposed Solution & Variants
            with st.expander("🛠️ Proposed Solution & Variants", expanded=False):
                st.session_state["final_prd"]["proposed_solution"] = st.text_area("Proposed Solution", value=st.session_state["final_prd"].get("proposed_solution", ""), height=120, key="edit_solution")
                variants = ensure_list(st.session_state["final_prd"].get("variants", []))
                if not variants:
                    variants = [{"control": "", "variation": "", "notes": ""}]
                v0 = variants[0]
                v0["control"] = st.text_area("Control", value=v0.get("control",""), height=80, key="edit_vctrl")
                v0["variation"] = st.text_area("Variation", value=v0.get("variation",""), height=80, key="edit_vvar")
                v0["notes"] = st.text_input("Notes", value=v0.get("notes",""), key="edit_vnotes")
                st.session_state["final_prd"]["variants"] = [v0]
                if st.button("♻️ Regenerate Solution & Variants", key="regen_variants"):
                    with st.spinner("Regenerating variants..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("final_prd", {}).get("hypotheses", [{}])[0]
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["proposed_solution"] = new_plan.get("proposed_solution", st.session_state["final_prd"].get("proposed_solution",""))
                                st.session_state["final_prd"]["variants"] = new_plan.get("variants", st.session_state["final_prd"].get("variants", []))
                                st.success("Solution & variants regenerated.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate variants: {e}")

            # Metrics
            with st.expander("📊 Metrics & Guardrails", expanded=False):
                metrics = ensure_list(st.session_state["final_prd"].get("metrics", []))
                if not metrics:
                    metrics = [{"name": "", "formula": "", "importance": "Primary"}]
                # render up to 3 metrics editable
                for i in range(len(metrics)):
                    metrics[i]["name"] = st.text_input(f"Metric {i+1} Name", value=metrics[i].get("name",""), key=f"edit_m_name_{i}")
                    metrics[i]["formula"] = st.text_input(f"Metric {i+1} Formula", value=metrics[i].get("formula",""), key=f"edit_m_for_{i}")
                    metrics[i]["importance"] = st.selectbox(f"Metric {i+1} Importance", ["Primary","Secondary"], index=0 if metrics[i].get("importance","Primary")=="Primary" else 1, key=f"edit_m_imp_{i}")
                st.session_state["final_prd"]["metrics"] = metrics
                if st.button("♻️ Regenerate Metrics", key="regen_metrics"):
                    with st.spinner("Regenerating metrics..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("final_prd", {}).get("hypotheses", [{}])[0]
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["metrics"] = new_plan.get("metrics", st.session_state["final_prd"].get("metrics", []))
                                st.session_state["final_prd"]["guardrail_metrics"] = new_plan.get("guardrail_metrics", st.session_state["final_prd"].get("guardrail_metrics", []))
                                st.success("Metrics regenerated.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate metrics: {e}")
            # Guardrail metrics
            with st.expander("🛡️ Guardrail Metrics", expanded=False):
                guardrails = ensure_list(st.session_state["final_prd"].get("guardrail_metrics", []))
                if not guardrails:
                    guardrails = [{"name": "", "direction": "Decrease", "threshold": ""}]
                for i in range(len(guardrails)):
                    guardrails[i]["name"] = st.text_input(f"Guardrail {i+1} Name", value=guardrails[i].get("name", ""), key=f"edit_g_name_{i}")
                    guardrails[i]["direction"] = st.selectbox(f"Guardrail {i+1} Direction", ["Increase", "Decrease", "No Change"], index=["Increase", "Decrease", "No Change"].index(guardrails[i].get("direction", "Decrease")), key=f"edit_g_dir_{i}")
                    guardrails[i]["threshold"] = st.text_input(f"Guardrail {i+1} Threshold", value=guardrails[i].get("threshold", ""), key=f"edit_g_thr_{i}")
                st.session_state["final_prd"]["guardrail_metrics"] = guardrails
                if st.button("♻️ Regenerate Guardrails", key="regen_guardrails"):
                    with st.spinner("Regenerating guardrails..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("final_prd", {}).get("hypotheses", [{}])[0]
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["guardrail_metrics"] = new_plan.get("guardrail_metrics", st.session_state["final_prd"].get("guardrail_metrics", []))
                                st.success("Guardrails regenerated.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate guardrails: {e}")

            # Risks & Assumptions
            with st.expander("⚠️ Risks & Assumptions", expanded=False):
                risks = ensure_list(st.session_state["final_prd"].get("risks_and_assumptions", []))
                if not risks:
                    risks = [{"risk": "", "severity": "Medium", "mitigation": ""}]
                for i in range(len(risks)):
                    risks[i]["risk"] = st.text_input(f"Risk {i+1}", value=risks[i].get("risk", ""), key=f"edit_risk_{i}")
                    risks[i]["severity"] = st.selectbox(f"Severity {i+1}", ["High", "Medium", "Low"], index=["High", "Medium", "Low"].index(risks[i].get("severity", "Medium")), key=f"edit_r_sev_{i}")
                    risks[i]["mitigation"] = st.text_input(f"Mitigation {i+1}", value=risks[i].get("mitigation", ""), key=f"edit_r_mit_{i}")
                st.session_state["final_prd"]["risks_and_assumptions"] = risks
                if st.button("♻️ Regenerate Risks", key="regen_risks"):
                    with st.spinner("Regenerating risks..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("final_prd", {}).get("hypotheses", [{}])[0]
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["risks_and_assumptions"] = new_plan.get("risks_and_assumptions", st.session_state["final_prd"].get("risks_and_assumptions", []))
                                st.success("Risks regenerated.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate risks: {e}")

            # Success & Learning Criteria
            with st.expander("📘 Success & Learning Criteria", expanded=False):
                sl = ensure_dict(st.session_state["final_prd"].get("success_learning_criteria", {}))
                sl["definition_of_success"] = st.text_input("Definition of Success", value=sl.get("definition_of_success", ""), key="edit_sl_def")
                sl["stopping_rules"] = st.text_input("Stopping Rules", value=sl.get("stopping_rules", ""), key="edit_sl_stop")
                sl["rollback_criteria"] = st.text_input("Rollback Criteria", value=sl.get("rollback_criteria", ""), key="edit_sl_roll")
                st.session_state["final_prd"]["success_learning_criteria"] = sl
                if st.button("♻️ Regenerate Success & Learning", key="regen_success_learning"):
                    with st.spinner("Regenerating success criteria..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("final_prd", {}).get("hypotheses", [{}])[0]
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["success_learning_criteria"] = new_plan.get("success_learning_criteria", st.session_state["final_prd"].get("success_learning_criteria", {}))
                                st.success("Success & Learning regenerated.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate success & learning: {e}")

            # Statistical Rationale
            with st.expander("📐 Statistical Rationale", expanded=False):
                st.session_state["final_prd"]["statistical_rationale"] = st.text_area("Statistical Rationale", value=st.session_state["final_prd"].get("statistical_rationale", ""), height=120, key="edit_stat")
                if st.button("♻️ Regenerate Statistical Rationale", key="regen_stats"):
                    with st.spinner("Regenerating statistical rationale..."):
                        try:
                            ctx = st.session_state.get("sidebar_context", {})
                            hyp = st.session_state.get("final_prd", {}).get("hypotheses", [{}])[0]
                            if PROMPT_ENGINE_AVAILABLE and _generate_experiment_plan:
                                raw_new = _generate_experiment_plan(ctx, hyp)
                                parsed_new = raw_new if isinstance(raw_new, dict) else extract_json_from_text(raw_new)
                                new_plan = sanitize_plan(parsed_new)
                                st.session_state["final_prd"]["statistical_rationale"] = new_plan.get("statistical_rationale", st.session_state["final_prd"].get("statistical_rationale", ""))
                                st.success("Statistical rationale regenerated.")
                            else:
                                st.warning("LLM not available.")
                        except Exception as e:
                            st.error(f"Failed to regenerate statistical rationale: {e}")

            # Save edits back
            st.session_state["final_prd"] = sanitize_plan(st.session_state["final_prd"])

            # Quick quality check
            if st.button("🔍 Quick AI Quality Check", key="qc_button"):
                with st.spinner("Running quality check..."):
                    try:
                        if PROMPT_ENGINE_AVAILABLE and _validate_experiment_plan:
                            issues = _validate_experiment_plan(st.session_state["final_prd"])
                            st.success("Quality check complete.")
                            st.info(json.dumps(issues, indent=2))
                        else:
                            st.info("No validation engine available. Consider reviewing metrics and guardrails manually.")
                    except Exception as e:
                        st.error(f"Quality check failed: {e}")

        # Preview tab
        with tab2:
            st.markdown("### Preview — Live document view")
            preview_md = prd_to_markdown(st.session_state["final_prd"])
            st.markdown(preview_md, unsafe_allow_html=True)

            # quick inline actions
            col_p1, col_p2, col_p3 = st.columns([1, 1, 1])
            with col_p1:
                if st.button("Export: Markdown", key="export_md"):
                    md_text = prd_to_markdown(st.session_state["final_prd"])
                    st.download_button("Download .md", data=md_text, file_name="experiment_prd.md", mime="text/markdown", key="dl_md_btn")
            with col_p2:
                if st.button("Export: JSON", key="export_json"):
                    json_str = json.dumps(st.session_state["final_prd"], indent=2)
                    st.download_button("Download .json", data=json_str, file_name="experiment_prd.json", mime="application/json", key="dl_json_btn")
            with col_p3:
                # PDF/DOCX export with availability checks
                if st.button("Export: PDF", key="export_pdf"):
                    if PDF_AVAILABLE:
                        try:
                            pdf_bytes = generate_pdf_bytes_from_prd_dict(st.session_state["final_prd"])
                            st.download_button("Download .pdf", data=pdf_bytes, file_name="experiment_prd.pdf", mime="application/pdf", key="dl_pdf_btn")
                        except Exception as e:
                            st.error(f"PDF generation failed: {e}")
                    else:
                        st.warning("PDF export not available (reportlab not installed).")
            # DOCX as separate button below
            if st.button("Export: DOCX", key="export_docx"):
                if DOCX_AVAILABLE:
                    try:
                        docx_bytes = generate_docx_bytes_from_plan(st.session_state["final_prd"])
                        st.download_button("Download .docx", data=docx_bytes, file_name="experiment_prd.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx_btn")
                    except Exception as e:
                        st.error(f"DOCX generation failed: {e}")
                else:
                    st.warning("DOCX export not available (python-docx not installed).")

        # End if plan exists (tabs)
        st.markdown("---")
        # Reset and minor actions
        colr1, colr2 = st.columns([1, 1])
        with colr1:
            if st.button("🔁 Reset PRD", key="reset_prd"):
                if "experiment_plan" in st.session_state:
                    del st.session_state["experiment_plan"]
                if "final_prd" in st.session_state:
                    del st.session_state["final_prd"]
                st.success("PRD reset. You can generate a new one.")
        with colr2:
            if st.button("🧾 Save to session", key="save_session"):
                st.session_state["experiment_plan"] = st.session_state.get("final_prd", {})
                st.success("Saved current PRD to session_state['experiment_plan'].")

    # Floating Tips Panel (rendered outside main flow so it overlaps)
    # Show if session_state flag
    try:
        if st.session_state.get("show_tips_panel", False):
            # Build context for tips
            ctx = st.session_state.get("sidebar_context", {})
            # Current step detection
            if "final_prd" in st.session_state:
                step = "prd"
            elif "chosen_hypothesis" in st.session_state:
                step = "hypothesis"
            else:
                step = "inputs"

            # Container styling via markdown (floating panel)
            panel_html = """
            <div class="floating-tips-panel">
              <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px">
                <div style="font-weight:700">💡 Tips</div>
                <div><button id="close_tips" style="border:none;background:#fff;cursor:pointer;font-size:14px">❌</button></div>
              </div>
            """
            # render existing tips if any
            tips = st.session_state.get("tips_list", [])
            if tips:
                for t in tips:
                    panel_html += f'<div class="tip-card">{t}</div>'
            else:
                panel_html += '<div class="muted">No tips yet. Click "Refresh Tips" to get tailored guidance.</div>'

            panel_html += """
              <div style="display:flex;gap:8px;margin-top:8px">
                <form action="#" method="post">
                  <input type="submit" id="refresh_tips" value="🔄 Refresh Tips" style="padding:8px 12px;border-radius:8px;border:1px solid #e2e8f0;background:#eef2ff;cursor:pointer;font-weight:600">
                </form>
              </div>
            </div>
            """

            st.markdown(panel_html, unsafe_allow_html=True)

            # Buttons can't be captured inside the raw HTML reliably; provide Streamlit controls under panel for actions.
            refresh = st.button("🔄 Refresh Tips", key="refresh_tips_streamlit")
            close = st.button("Close Tips", key="close_tips_streamlit")
            if refresh:
                with st.spinner("Refreshing tips..."):
                    try:
                        tips_out = generate_tips(ctx, step)
                        # normalize to strings (limit 5)
                        normalized = [str(x).strip() for x in (tips_out or [])][:5]
                        st.session_state["tips_list"] = normalized
                        st.success("Tips updated.")
                    except Exception as e:
                        st.error(f"Failed to get tips: {e}")
            if close:
                st.session_state["show_tips_panel"] = False

    except Exception:
        # If anything with tips rendering fails, fail silently (UX only)
        pass

    # Floating button to open tips (bottom-right)
    floating_html = """
    <div class="floating-tips-button">
      <button onclick="window.parent.document.querySelector('button[kind=\"primary\"]').click();" style="background:#6366f1;color:#fff;border:none;padding:12px;border-radius:999px;box-shadow:0 6px 18px rgba(99,102,241,0.3);cursor:pointer;font-weight:700">
        💡 Tips
      </button>
    </div>
    """
    # We render a static floating button; functionality is via the header toggles and streamlit buttons.
    st.markdown(floating_html, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
