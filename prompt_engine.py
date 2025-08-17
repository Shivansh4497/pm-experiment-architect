import os
import json
import re
import io
from typing import Dict, Any, List

import openai
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document

# ============ LLM Client Setup ============
GROQ_AVAILABLE = False
_client = None
try:
    # Import guarded so module still loads if Groq not installed
    from groq import Groq  # type: ignore
    GROQ_AVAILABLE = True
    # instantiate safely (expect user to set GROQ_API_KEY in env)
    try:
        _client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
    except Exception:
        # keep _client None if init failed
        _client = None
except Exception:
    GROQ_AVAILABLE = False
    _client = None

# ============ Prompt Templates ============
PROMPTS = {
    "hypotheses": """
You are an expert Product Manager. Based on the user inputs, generate 3 highly relevant A/B test hypotheses. 
Each hypothesis must be structured JSON with the following fields:
- hypothesis: a single clear testable statement
- rationale: why this change could work, linked to business goal + user persona
- example_implementation: one concrete way to run this test
- behavioral_basis: psychological / behavioral science principle backing this

User Inputs:
Business Goal: {business_goal}
Product Type: {product_type}
Target Persona: {user_persona}
Key Metric: {key_metric}
Current Value: {current_value}
Target Value: {target_value}

Return JSON only in this structure:
{{
  "hypotheses": [
    {{
      "hypothesis": "...",
      "rationale": "...",
      "example_implementation": "...",
      "behavioral_basis": "..."
    }}
  ]
}}
""",

    "prd": """
You are a senior product manager. Generate a complete A/B Test PRD in structured JSON.

Inputs:
Business Goal: {business_goal}
Product Type: {product_type}
Target Persona: {user_persona}
Key Metric: {key_metric}
Current Value: {current_value}
Target Value: {target_value}
Hypothesis: {hypothesis}

Output JSON must have:
{{
  "experiment_title": "...",
  "problem_statement": "...",
  "hypothesis": "...",
  "proposed_solution": {{
    "control": "...",
    "variant": "..."
  }},
  "metrics": {{
    "primary": "...",
    "secondary": ["..."],
    "guardrails": ["..."]
  }},
  "design": {{
    "sample_size_per_variant": 0,
    "total_sample_size": 0,
    "confidence_level": 95.0,
    "statistical_power": 80.0,
    "mde": 0.05,
    "dau_coverage": 0.1
  }},
  "risks": [
    {{ "risk": "...", "mitigation": "..." }}
  ],
  "success_criteria": "...",
  "learning_criteria": "..."
}}
""",

    "tips": """
You are acting as a mentor for a product manager writing A/B test PRDs.
Provide 3-5 short, practical tips for the current step: {step}.
Base tips on the context:
{context}

Keep tips concrete, no fluff. Format:
[
  "💡 Tip 1...",
  "✅ Tip 2...",
  "⚠️ Tip 3..."
]
"""
}


# ============ Utilities ============

def extract_json_from_text(text: str) -> dict:
    """Extract JSON from LLM output safely."""
    if not text:
        return {}
    try:
        # Strip Markdown formatting if present
        cleaned = re.sub(r"^```json|```$", "", text.strip(), flags=re.MULTILINE).strip()
        return json.loads(cleaned)
    except Exception:
        try:
            return json.loads(text)
        except Exception:
            return {}

def safe_call_llm(prompt: str, model: str = "gpt-4o-mini", temperature: float = 0.7) -> str:
    """Wrapper for OpenAI API with fallback on failure."""
    try:
        resp = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
        )
        return resp.choices[0].message["content"]
    except Exception as e:
        print(f"LLM call failed: {e}")
        return ""
# ============ Hypothesis Generation ============

def generate_hypotheses(context: dict) -> List[Dict[str, str]]:
    """Generate 3 hypotheses based on user context, always return list of dicts."""
    prompt = PROMPTS["hypotheses"].format(**context)
    raw = safe_call_llm(prompt)
    parsed = extract_json_from_text(raw)

    hyps = []
    if parsed and "hypotheses" in parsed:
        for h in parsed["hypotheses"]:
            if isinstance(h, dict):
                hyps.append({
                    "hypothesis": h.get("hypothesis", "").strip(),
                    "rationale": h.get("rationale", "").strip(),
                    "example_implementation": h.get("example_implementation", "").strip(),
                    "behavioral_basis": h.get("behavioral_basis", "").strip(),
                })
            else:
                hyps.append({
                    "hypothesis": str(h).strip(),
                    "rationale": "",
                    "example_implementation": "",
                    "behavioral_basis": ""
                })
    return hyps


def expand_hypothesis_with_details(hypothesis: str, context: dict) -> dict:
    """Enrich a chosen hypothesis with rationale, example, basis."""
    prompt = f"""
You are expanding a hypothesis into full detail for an A/B test PRD.

Context:
Business Goal: {context.get("business_goal")}
Product Type: {context.get("product_type")}
User Persona: {context.get("user_persona")}
Key Metric: {context.get("key_metric")}

Hypothesis: {hypothesis}

Return JSON with:
{{
  "hypothesis": "...",
  "rationale": "...",
  "example_implementation": "...",
  "behavioral_basis": "..."
}}
"""
    raw = safe_call_llm(prompt)
    parsed = extract_json_from_text(raw)
    if parsed:
        return {
            "hypothesis": parsed.get("hypothesis", hypothesis),
            "rationale": parsed.get("rationale", ""),
            "example_implementation": parsed.get("example_implementation", ""),
            "behavioral_basis": parsed.get("behavioral_basis", "")
        }
    return {"hypothesis": hypothesis, "rationale": "", "example_implementation": "", "behavioral_basis": ""}


# ============ PRD Generation ============

def generate_prd(context: dict, hypothesis: dict) -> dict:
    """Generate a full PRD from context + hypothesis."""
    prompt = PROMPTS["prd"].format(
        business_goal=context.get("business_goal", ""),
        product_type=context.get("product_type", ""),
        user_persona=context.get("user_persona", ""),
        key_metric=context.get("key_metric", ""),
        current_value=context.get("current_value", ""),
        target_value=context.get("target_value", ""),
        hypothesis=hypothesis.get("hypothesis", "")
    )
    raw = safe_call_llm(prompt, temperature=0.5)
    parsed = extract_json_from_text(raw)
    return sanitize_experiment_plan(parsed)


def generate_experiment_plan(context: dict, hypothesis: dict) -> dict:
    """Alias for backward compatibility with main.py."""
    return generate_prd(context, hypothesis)


# ============ Sanitizers ============

def sanitize_experiment_plan(raw_plan: dict) -> dict:
    """Ensure all expected keys exist with safe defaults."""
    if not isinstance(raw_plan, dict):
        raw_plan = {}

    return {
        "experiment_title": raw_plan.get("experiment_title", "Untitled Experiment"),
        "problem_statement": raw_plan.get("problem_statement", ""),
        "hypothesis": raw_plan.get("hypothesis", ""),
        "proposed_solution": {
            "control": raw_plan.get("proposed_solution", {}).get("control", ""),
            "variant": raw_plan.get("proposed_solution", {}).get("variant", "")
        },
        "metrics": {
            "primary": raw_plan.get("metrics", {}).get("primary", ""),
            "secondary": raw_plan.get("metrics", {}).get("secondary", []),
            "guardrails": raw_plan.get("metrics", {}).get("guardrails", [])
        },
        "design": {
            "sample_size_per_variant": raw_plan.get("design", {}).get("sample_size_per_variant", 0),
            "total_sample_size": raw_plan.get("design", {}).get("total_sample_size", 0),
            "confidence_level": raw_plan.get("design", {}).get("confidence_level", 95.0),
            "statistical_power": raw_plan.get("design", {}).get("statistical_power", 80.0),
            "mde": raw_plan.get("design", {}).get("mde", 0.05),
            "dau_coverage": raw_plan.get("design", {}).get("dau_coverage", 0.1),
        },
        "risks": raw_plan.get("risks", []),
        "success_criteria": raw_plan.get("success_criteria", ""),
        "learning_criteria": raw_plan.get("learning_criteria", "")
    }


def sanitize_plan(plan: dict) -> dict:
    """Alias for compatibility, same as sanitize_experiment_plan."""
    return sanitize_experiment_plan(plan)
# ============ Export Utilities ============

def generate_pdf_bytes_from_prd_dict(plan: dict) -> bytes:
    """Generate PDF export from PRD dict."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    def add_section(title: str, content: str):
        story.append(Paragraph(f"<b>{title}</b>", styles["Heading3"]))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(content or "-", styles["Normal"]))
        story.append(Spacer(1, 0.2 * inch))

    add_section("Experiment Title", plan.get("experiment_title", ""))
    add_section("Problem Statement", plan.get("problem_statement", ""))
    add_section("Hypothesis", plan.get("hypothesis", ""))

    ps = plan.get("proposed_solution", {})
    add_section("Proposed Solution - Control", ps.get("control", ""))
    add_section("Proposed Solution - Variant", ps.get("variant", ""))

    metrics = plan.get("metrics", {})
    add_section("Primary Metric", metrics.get("primary", ""))
    add_section("Secondary Metrics", ", ".join(metrics.get("secondary", [])))
    add_section("Guardrail Metrics", ", ".join(metrics.get("guardrails", [])))

    design = plan.get("design", {})
    add_section("Design", json.dumps(design, indent=2))

    risks = plan.get("risks", [])
    risks_text = "<br/>".join([f"- {r.get('risk')}: {r.get('mitigation')}" for r in risks])
    add_section("Risks & Mitigation", risks_text)

    add_section("Success Criteria", plan.get("success_criteria", ""))
    add_section("Learning Criteria", plan.get("learning_criteria", ""))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


def generate_docx_bytes_from_prd_dict(plan: dict) -> bytes:
    """Generate DOCX export from PRD dict."""
    buffer = io.BytesIO()
    doc = Document()

    def add_section(title: str, content: str):
        doc.add_heading(title, level=2)
        doc.add_paragraph(content or "-")

    add_section("Experiment Title", plan.get("experiment_title", ""))
    add_section("Problem Statement", plan.get("problem_statement", ""))
    add_section("Hypothesis", plan.get("hypothesis", ""))

    ps = plan.get("proposed_solution", {})
    add_section("Proposed Solution - Control", ps.get("control", ""))
    add_section("Proposed Solution - Variant", ps.get("variant", ""))

    metrics = plan.get("metrics", {})
    add_section("Primary Metric", metrics.get("primary", ""))
    add_section("Secondary Metrics", ", ".join(metrics.get("secondary", [])))
    add_section("Guardrail Metrics", ", ".join(metrics.get("guardrails", [])))

    design = plan.get("design", {})
    add_section("Design", json.dumps(design, indent=2))

    risks = plan.get("risks", [])
    for r in risks:
        doc.add_paragraph(f"- {r.get('risk')}: {r.get('mitigation')}")

    add_section("Success Criteria", plan.get("success_criteria", ""))
    add_section("Learning Criteria", plan.get("learning_criteria", ""))

    doc.save(buffer)
    return buffer.getvalue()


# ============ Dynamic Tips ============

def generate_tips(step: str, context: dict) -> List[str]:
    """Generate contextual PM tips for a given step."""
    prompt = PROMPTS["tips"].format(step=step, context=json.dumps(context, indent=2))
    raw = safe_call_llm(prompt, temperature=0.6)
    parsed = extract_json_from_text(raw)

    # Fallback: return as list of strings
    if isinstance(parsed, list):
        return parsed
    elif isinstance(parsed, dict) and "tips" in parsed:
        return parsed["tips"]
    elif isinstance(raw, str):
        # crude split fallback
        return [line.strip() for line in raw.split("\n") if line.strip()]
    return ["⚠️ No tips generated. Try refreshing."]
